import docx, time, os
import win32com.client as client

def convert_to_pdf(doc):
        word = client.DispatchEx("Word.Application")
        new_name = doc.replace(".docx", r".pdf")
        worddoc = word.Documents.Open(doc)
        worddoc.SaveAs(new_name, FileFormat = 17)
        worddoc.Close()

# Open docx file
d=docx.Document('consent.docx')

# Get student and parents name
nome_estudante=input("Student Name: ")
d.paragraphs[3].runs[2].text=nome_estudante
parent1=input("Parent 1: ")
d.paragraphs[7].runs[0].text=parent1
parent2=input("Parent 2: ")
d.paragraphs[9].runs[0].text=parent2

# Date
data1=time.strftime("	%d/%m/%Y")
d.paragraphs[7].runs[1].text=data1
d.paragraphs[9].runs[1].text=data1

seguraloop=1
#Consent
while seguraloop==1:
    consent=input("Do you consent for evaluation (yes/no): ")
    consent.casefold()
    if consent=='yes':
        d.paragraphs[4].runs[0].text='    X    '
        d.paragraphs[5].runs[0].text='          '
        d.paragraphs[4].runs[0].underline=True
        d.paragraphs[5].runs[0].underline=True
        seguraloop=0
    elif consent=='no':
        d.paragraphs[4].runs[0].text='          '
        d.paragraphs[5].runs[0].text='    X    '
        d.paragraphs[4].runs[0].underline=True
        d.paragraphs[5].runs[0].underline=True
        seguraloop=0
    else:
        print('\nInvalid answer, please answer yes or no:')
        seguraloop=1

# Save as DOCX and PDF
d.save('output.docx')
convert_to_pdf(os.getcwd()+'\\output.docx')
print('\nFinalizado !')